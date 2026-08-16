import os
from docx import Document

def create_docx():
    doc = Document()
    doc.add_heading('Deep Reinforcement Learning for Proactive Network Traffic Congestion Control and Load Balancing', 0)
    
    doc.add_heading('Authors', level=1)
    doc.add_paragraph('Students of University of Ghana')
    doc.add_paragraph('College of Basic and Applied Sciences')
    doc.add_paragraph('Department of Computer Science')
    doc.add_paragraph('Master of Science Computer Science 2025/2026')
    
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('Modern ISP and enterprise networks suffer from transient micro-bursts and rapidly fluctuating traffic demands. Traditional traffic engineering approaches such as Equal-Cost Multi-Path (ECMP) rely on static hash-based packet distribution. While mathematically simple, these methods are blind to real-time network states and often hash heavy "elephant" flows onto the same physical links, causing severe bottlenecks and buffer bloat. This paper evaluates Deep Reinforcement Learning (DRL) algorithms—specifically Deep Q-Network (DQN) and Proximal Policy Optimization (PPO)—for proactive load balancing across wide-area topologies. Our agents are trained to dynamically shift flows across candidate paths to minimize queue congestion and maximize throughput. Tested on unseen topologies from the Internet Topology Zoo, PPO and DQN substantially outperformed the ECMP baseline during highly variable traffic bursts.')
    
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('Network traffic engineering aims to optimize performance by dynamically routing data. ECMP provides static multi-path routing but fails under asymmetric loads. We propose using DRL to proactively manage flows.')
    
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph('We modeled the network as a Markov Decision Process (MDP) using Gymnasium. The state space consists of link utilization, queue occupancy, and flow demands. The action space allows the agent to shift flows between precomputed k-shortest paths. We trained DQN and PPO agents over 5 random seeds using Stable-Baselines3.')
    
    doc.add_heading('3. Results', level=1)
    doc.add_paragraph('Our experiments show that PPO achieves highly stable convergence compared to DQN. Under burst traffic conditions, DRL agents achieved lower congestion penalties than ECMP. Detailed tables and figures are available in the supplementary material.')
    
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph('DRL provides a viable alternative to static hashing algorithms for load balancing in dynamic networks. PPO demonstrated the most robust generalization to unseen topologies.')
    
    doc.save('final_report.docx')
    print("final_report.docx created successfully.")

if __name__ == '__main__':
    create_docx()
